import os
import csv
import pandas as pd
from jobspy import scrape_jobs
from openai import OpenAI
import json
from docx import Document
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import datetime
import os
from openai import OpenAI
from dotenv import load_dotenv  # <--- This library reads the .env file

# 1. Load the secrets from the .env file
load_dotenv()

# 2. Get the key from the loaded environment
api_key_from_env = os.getenv("OPENAI_API_KEY")

# 3. specific check to ensure it loaded
if not api_key_from_env:
    print("❌ ERROR: Could not find 'OPENAI_API_KEY'. Please check your .env file.")
    exit()

# 4. Initialize the client using the variable
client = OpenAI(api_key=api_key_from_env)

# ==========================================
# REST OF YOUR CODE BELOW...
# ==========================================
# ==========================================
# CONFIGURATION
# ==========================================
# API Keys (Now loaded securely)
API_KEY = os.getenv("OPENAI_API_KEY")
EMAIL_SENDER = os.getenv("EMAIL_SENDER")
EMAIL_PASSWORD = os.getenv("EMAIL_PASSWORD") # App Password
EMAIL_RECEIVER = os.getenv("EMAIL_RECEIVER")

# SEARCH PARAMS
SEARCH_TERM = (
    'title:(("AI" OR "Artificial Intelligence" OR "Machine Learning" OR "ML" OR '
    '"Deep Learning" OR "Computer Vision" OR "NLP" OR "GenAI" OR "Generative AI" OR '
    '"LLM" OR "Applied AI" OR "MLOps") AND "Intern")'
)

LOCATION = "United States" 
JOBS_TO_SCRAPE = 20 

client = OpenAI(api_key=API_KEY)

# ==========================================
# YOUR RESUME
# ==========================================
MY_RESUME = """
EDGAR PFUMA
Intern 2026: Software Engineer - Semiconductor
Dallas–Fort Worth, TX • edgarpfuma@hotmail.com • (945) 7304333 • LinkedIn: www.linkedin.com/in/edgar-pfuma
(Rest of resume hidden for brevity, assumes it is loaded...)
"""

# ==========================================
# SYSTEM PROMPT
# ==========================================
SYSTEM_PROMPT = """
Role: Act as a Strategic Career Coach & Technical Recruiter specializing in US Immigration for Tech.
Context: Candidate is MS AI student (Graduating May 2026). Currently on CPT.
Task: Analyze Job Description.
Return JSON:
1. "verdict": "APPLY (Long-Term)", "APPLY (Resume Builder)", or "STOP".
2. "reasoning": Brief explanation.
3. "compatibility_score": 0-100.
4. "optimized_summary": Rewritten professional summary.
5. "tailored_bullets": Array of 3 rewritten bullets.
6. "cover_letter_hook": 2-sentence LinkedIn message.
7. "company_name": Extract company name.
Output Format: JSON ONLY.
"""

# ==========================================
# FUNCTIONS
# ==========================================

def get_jobs_online():
    print(f"🕵️  Scraping {JOBS_TO_SCRAPE} jobs for '{SEARCH_TERM}' in {LOCATION}...")
    try:
        jobs = scrape_jobs(
            # 1. ADD THESE NEW SITES
            site_name=["indeed", "glassdoor", "zip_recruiter", "google"],
            
            search_term=SEARCH_TERM,
            location=LOCATION,
            results_wanted=JOBS_TO_SCRAPE,
            hours_old=72,
            country_indeed='USA',
            
            # 2. GOOGLE SPECIFIC SETTING
            # Google needs a slightly different search query format sometimes
            google_search_term=f"{SEARCH_TERM} jobs in {LOCATION}"
        )
        
        # --- DEDUPLICATION LOGIC ---
        if not jobs.empty:
            initial_count = len(jobs)
            
            # Create unique ID from Company + Title
            jobs['dedup_id'] = jobs['company'].astype(str).str.upper() + "_" + jobs['title'].astype(str).str.upper()
            jobs = jobs.drop_duplicates(subset=['dedup_id'], keep='first')
            
            final_count = len(jobs)
            print(f"✅ Found {initial_count} jobs (Removed {initial_count - final_count} duplicates). Final count: {final_count}")
        
        return jobs

    except Exception as e:
        print(f"❌ Scraping failed: {e}")
        return pd.DataFrame()

def pre_filter_job(job_data):
    """
    LIGHTWEIGHT FILTER:
    1. Title Check: Must match our "Safe" keywords.
    2. Citizenship Check: Filters out clearance jobs.
    3. NO SKILL CHECK: We apply to everything relevant.
    """
    if job_data.empty: return False

    title = str(job_data.get('title', '')).lower()
    description = str(job_data.get('description', '')).lower()
    
    # --- CHECK 1: BAD TITLES (Instant Reject) ---
    bad_keywords = ["sales", "marketing", "finance", "accounting", "hr", "maintenance", 
                    "technician", "mechanic", "audit", "legal", "driver", "nurse", "warehouse",
                    "customer service", "executive assistant", "recruiter"]
    if any(k in title for k in bad_keywords):
        print(f"   (Skipping: Non-Tech Title '{job_data.get('title')}')")
        return False

    # --- CHECK 2: CITIZENSHIP (Instant Reject) ---
    citizenship_blocklist = ["us citizen only", "u.s. citizen only", "security clearance", 
                             "active secret", "dod clearance", "citizenship required"]
    for word in citizenship_blocklist:
        if word in description:
            print(f"   (Skipping: Citizenship keyword '{word}')")
            return False
            
    # If it passed the "Bad Title" and "Citizenship" checks, we ACCEPT it.
    return True

def analyze_with_ai(job_description):
    # --- NEW SAFETY CHECK ---
    # If the description is "nan" or too short, don't waste money asking AI.
    if not job_description or str(job_description).lower() == 'nan' or len(str(job_description)) < 50:
        return {
            "verdict": "STOP", 
            "reasoning": "Scraper failed to get job description text (LinkedIn anti-bot likely active)."
        }

    try:
        response = client.chat.completions.create(
            # ... existing code ...
            model="gpt-4-turbo",
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": f"RESUME:\n{MY_RESUME}\n\nJOB DESCRIPTION:\n{job_description}"}
            ],
            response_format={"type": "json_object"},
            temperature=0.7
        )
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        return {"error": str(e)}

def save_application_kit(analysis, job_data):
    """Generates the Word Doc and returns the filename"""
    company = analysis.get('company_name', job_data.get('company', 'Company'))
    clean_company = "".join(x for x in str(company) if x.isalnum() or x in "._- ")
    filename = f"App_Kit_{clean_company}.docx"
    
    try:
        doc = Document()
        doc.add_heading(f'Application Kit: {company}', 0)
        doc.add_paragraph(f"Job Title: {job_data.get('title')}")
        doc.add_paragraph(f"Source: {job_data.get('site', 'Unknown')}")
        doc.add_paragraph(f"URL: {job_data.get('job_url_direct')}")
        doc.add_paragraph("-" * 20)
        doc.add_heading('Verdict', level=1)
        doc.add_paragraph(f"{analysis.get('verdict')} - {analysis.get('reasoning')}")
        doc.add_heading('LinkedIn Message', level=1)
        p = doc.add_paragraph(analysis.get('cover_letter_hook'))
        p.style = 'Intense Quote'
        doc.add_heading('Optimized Summary', level=1)
        doc.add_paragraph(analysis.get('optimized_summary'))
        doc.add_heading('Tailored Bullets', level=1)
        for bullet in analysis.get('tailored_bullets', []):
            doc.add_paragraph(bullet, style='List Bullet')
        
        doc.save(filename)
        print(f"   ✅ Saved: {filename}")
        return filename  # Return the filename so we can email it later!
    except Exception as e:
        print(f"   ❌ Error saving file: {e}")
        return None

def send_email_report(successful_jobs):
    """Sends one email with all attachments"""
    if not successful_jobs:
        print("📭 No jobs to email today.")
        return

    print(f"\n📧 Sending email report with {len(successful_jobs)} attachments...")
    
    msg = MIMEMultipart()
    msg['From'] = EMAIL_SENDER
    msg['To'] = EMAIL_RECEIVER
    msg['Subject'] = f"🚀 AI Job Report: {len(successful_jobs)} New Opportunities ({datetime.date.today()})"

    body = "Here are your tailored application kits for today:\n\n"
    for job in successful_jobs:
        body += f"✅ {job['title']} @ {job['company']}\n   Strategy: {job['strategy']}\n   URL: {job['url']}\n\n"
    
    body += "\nGood luck!\n- Your AI Agent"
    msg.attach(MIMEText(body, 'plain'))

    # Attach Files
    for job in successful_jobs:
        filename = job['file_path']
        if filename and os.path.exists(filename):
            try:
                with open(filename, "rb") as attachment:
                    part = MIMEBase("application", "octet-stream")
                    part.set_payload(attachment.read())
                encoders.encode_base64(part)
                part.add_header("Content-Disposition", f"attachment; filename= {filename}")
                msg.attach(part)
            except Exception as e:
                print(f"   Could not attach {filename}: {e}")

    # Send
    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(EMAIL_SENDER, EMAIL_PASSWORD)
        text = msg.as_string()
        server.sendmail(EMAIL_SENDER, EMAIL_RECEIVER, text)
        server.quit()
        print("✅ EMAIL SENT SUCCESSFULLY!")
    except Exception as e:
        print(f"❌ Failed to send email: {e}")

# ==========================================
# MAIN EXECUTION
# ==========================================
if __name__ == "__main__":
    jobs_df = get_jobs_online()
    successful_applications = [] # List to store jobs for email

    if not jobs_df.empty:
        output_folder = "Todays_Job_Batch"
        os.makedirs(output_folder, exist_ok=True)
        os.chdir(output_folder)
        
        print("\n🚀 Starting AI Analysis Pipeline...\n")
        
        for index, row in jobs_df.iterrows():
            title = row.get('title', 'Unknown Role')
            company = row.get('company', 'Unknown Company')
            description = row.get('description', '')
            site = row.get('site', 'Unknown')
            
            print(f"🔎 Checking: {title} at {company} via {site}...")
            
            if pre_filter_job(row):
                analysis = analyze_with_ai(description)
                
                if "error" in analysis:
                    print(f"   ⚠️ AI Error: {analysis['error']}")
                    continue
                
                verdict = analysis.get('verdict', 'STOP')
                
                if "APPLY" in verdict:
                    # Generate File
                    filename = save_application_kit(analysis, row)
                    
                    # Add to Email List
                    if filename:
                        successful_applications.append({
                            'title': title,
                            'company': company,
                            'strategy': analysis.get('reasoning'),
                            'url': row.get('job_url_direct'),
                            'file_path': filename
                        })
                else:
                    print(f"   🛑 Verdict: STOP ({analysis.get('reasoning')})")
            
        print("\n🏁 Batch complete.")
        
        # Send Email
        if len(successful_applications) > 0:
            send_email_report(successful_applications)
        else:
            print("No valid applications found today, so no email sent.")
            
    else:
        print("No jobs found.")